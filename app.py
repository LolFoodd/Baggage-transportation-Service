import sqlite3
from flask import send_from_directory, send_file
from werkzeug.exceptions import abort
import os

from io import BytesIO
import pandas as pd

from flask import Flask, render_template, request, redirect, url_for, flash, session, abort
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user

import pprint
from docx import Document

import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

import xlsxwriter


app = Flask(__name__)
app.config['SECRET_KEY'] = b'my)secret)key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

UPLOAD_FOLDER = 'contracts'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# User model
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(150), unique=True, nullable=False)
    password = db.Column(db.String(150), nullable=False)
    role = db.Column(db.String(50), nullable=False)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


with app.app_context():
    db.create_all()


# Decorator for role-based access
def admin_required(f):
    def wrap(*args, **kwargs):
        if current_user.role != 'admin':
            flash('Доступ запрещен.')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    wrap.__name__ = f.__name__
    return wrap

# Create all tables
with app.app_context():
    db.create_all()
    # Add users to the database
    admin = User.query.filter_by(username='admin').first()
    if not admin:
        admin = User(username='admin', password='password_admin', role='admin')
        db.session.add(admin)
    user = User.query.filter_by(username='user').first()
    if not user:
        user = User(username='user', password='password_user', role='user')
        db.session.add(user)
    db.session.commit()



@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username, password=password).first()
        if user:
            login_user(user)
            return redirect(url_for('index'))
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))



def get_db_connection():
    conn = sqlite3.connect('database1.db')
    conn.row_factory = sqlite3.Row
    return conn

@app.route('/')
@login_required
def index():
    return redirect("/contracts")


# Контракты
""" Страница-список - получение всех контрактов """
@app.route('/contracts')
@login_required
def contracts():
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM contracts, things, clients
    WHERE contracts.thing_id = things.id_thing and contracts.client_id = clients.id_client
    """).fetchall()
    conn.close()
    return render_template('contracts.html', contracts=pos)




""" Получение одного контракта из БД """
def get_contract(item_id):
    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM contracts, things, clients, employees
    WHERE contracts.thing_id = things.id_thing and contracts.employee_id = employees.id_employee and contracts.client_id = clients.id_client and contracts.id_contract = ?
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item




""" Страница-карточка - 1 контракт """
@app.route('/contract/<int:contract_id>')
def contract(contract_id):
    pos = get_contract(contract_id)
    return render_template('contract.html', contract=pos)




""" Страница-добавления нового контракта """
@app.route('/new_contract', methods=('GET', 'POST'))
@login_required
@admin_required
def new_contract():

    if request.method == 'POST':
        # добавление нового контракта в БД псоле заполнения формы
        try:
            number = request.form['number']
            date = request.form['date']
            start_price = int(request.form['start_price'])
            discount = int(request.form['discount'])
            deal_status = 0
            finish_price = int(start_price * ((100 - discount) / 100))
            address_start = request.form['address_start']
            address_finish = request.form['address_finish']

            thing_info = request.form['thing'].split(',')
            thing_id = int(thing_info[0])
            client_id = int(thing_info[1])

            employee_id = int(request.form.get('employee'))
        except ValueError:
            flash('Некорректные значения')
            client_id = 0
        if not (client_id > 0 and thing_id > 0 and employee_id > 0):
            flash('Не все поля заполнены')
        else:
            if not (number and date and start_price and address_start and address_finish):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute("INSERT INTO 'contracts' ('number', 'date', 'start_price', 'discount', 'deal_status', 'finish_price','address_start', 'address_finish','thing_id', 'client_id', 'employee_id')  VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                                                       (number, date, start_price, discount, deal_status, finish_price, address_start, address_finish, thing_id, client_id, employee_id))
                conn.commit()
                new_contract_id = cursor.lastrowid
                print(new_contract_id)
                conn.close()
                return redirect(f'/contract/{new_contract_id}')

    # отрисовка формы
    conn = get_db_connection()
    pos1 = conn.execute("""SELECT * FROM clients""").fetchall()
    pos2 = conn.execute("""SELECT * FROM things WHERE id_thing NOT IN (SELECT thing_id FROM contracts)""").fetchall()
    pos3 = conn.execute("""SELECT * FROM employees""").fetchall()
    conn.close()
    return render_template('new_contract.html', clients=pos1, things=pos2, employees=pos3)








@app.route('/generate_contract', methods=('GET', 'POST'))
def generate_contract():
    """ Страница генерации договора """

    # переменные шаблона
    id = int(request.args.get('id_contract'))
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM contracts, clients, employees
            WHERE contracts.client_id = clients.id_client and 
                        contracts.employee_id = employees.id_employee and
                        contracts.id_contract = ?
                        """, (id,)).fetchone()
    conn.close()
    contract_params = {
            'EMPLOYEE_ATTORNEY_POWER': 'основание действий сотрудника',
            'CLIENT_BIRTHDATE': 'дата рождения клиента',
            'CONTRACT_CITY': 'город заключения конракта',
            'CLIENT_BIRTHPLACE': 'место рождения клиента',
            'CLIENT_PASSPORT_DEPCODE': 'код подразделения, выдавшего паспорт клиента',
            'CLIENT_RS': 'расчетный счет клиента',
            'CLIENT_BANK': 'Банк клиента',
            'CLIENT_BIK': 'БИК клиента'}

    contract_params_auto = {
            'CONTRACT_NUMBER': ['номер договора', pos['number']],
            'CONTRACT_DATE': ['дата подписания договора', pos['date']],
            'CLIENT_FULLNAME': ['ФИО клиента', pos['name']],
            'CLIENT_PASSPORT_NUMBER': ['серия и номер паспорта клиента', pos['passport']],
            'EMPLOYEE_POSITION': ['должность сотрудника', pos['position']],
            'EMPLOYEE_FULLNAME': ['ФИО сотрудника', pos[18]],
            'ADDRESS_START': ['стартовый адрес ', pos['address_start']],
            'ADDRESS_FINISH': ['адрес доставки', pos['address_finish']],
            'CONTRACT_SUM': ['сумма услуг', pos['finish_price']]}

    if request.method == 'POST':
        # создание нового документа
        result_params =  request.form.to_dict()
        create_contract(id, result_params)
        return redirect(f'/contract/{id}')

    # скачивание файла, если он заполнен
    filename = f"договор {pos['number']} от {pos['date']}.docx"
    if os.path.exists(os.path.join('contracts', filename)):
        return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)
    else:
        # отрисовка формы заполнения
        flash('Договор не сформирован, заполните его')
        return render_template('generate_contract.html',
                               contract=pos, contract_params=contract_params, auto_params=contract_params_auto)


def create_contract(id, contract_params):
    """ Создание нового документа по шаблону """

    template = os.path.join('contracts', 'contract_template.docx')
    result = os.path.join('contracts', f"договор {contract_params['CONTRACT_NUMBER']} от {contract_params['CONTRACT_DATE']}.docx")

    template_doc = Document(template)
    for key, value in contract_params.items():
        for paragraph in template_doc.paragraphs:
            replace_text(paragraph, f'=={key}==', value)
        for table in template_doc.tables:
            replace_text_in_tables(table, key, value)
    template_doc.save(result)


def replace_text(paragraph, key, value):
  """ Работа docx - заполнение параграфов """

  if key in paragraph.text:
    paragraph.text = paragraph.text.replace(key, value)


def replace_text_in_tables(table, key, value):
  """ Работа docx - заполнение таблиц """
  for row in table.rows:
    for cell in row.cells:
      if key in cell.text:
        cell.text = cell.text.replace(key, value)



""" Страница-редактирования контракта """


@app.route('/edit_contract/<int:contract_id>', methods=('GET', 'POST'))
def edit_contract(contract_id):
    contract = get_contract(contract_id)

    if request.method == 'POST':
        try:
            number = request.form['number']
            date = request.form['date']
            start_price = int(request.form['start_price'])
            discount = int(request.form['discount'])
            deal_status = int(request.form.get('deal_status', 0))
            finish_price = int(start_price * ((100 - discount) / 100))
            address_start = request.form['address_start']
            address_finish = request.form['address_finish']

            thing_info = request.form.get('thing', '')
            if thing_info:
                thing_info = thing_info.split(',')
                thing_id = int(thing_info[0])
                client_id = int(thing_info[1])
            else:
                thing_id = contract['thing_id']
                client_id = contract['client_id']

            employee_id = int(request.form.get('employee'))
        except ValueError:
            flash('Некорректные значения')
            return render_template('edit_contract.html', contract=contract, clients=clients, things=things,
                                   employees=employees)

        if not (client_id > 0 and thing_id > 0 and employee_id > 0):
            flash('Не все поля заполнены')
        else:
            if not (number and date and start_price and address_start and address_finish):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute("""
                    UPDATE contracts
                    SET number = ?, date = ?, start_price = ?, discount = ?, deal_status = ?, finish_price = ?, 
                        address_start = ?, address_finish = ?, thing_id = ?, client_id = ?, employee_id = ?
                    WHERE id_contract = ?
                """, (
                number, date, start_price, discount, deal_status, finish_price, address_start, address_finish, thing_id,
                client_id, employee_id, contract_id))
                conn.commit()
                conn.close()
                return redirect(f'/contract/{contract_id}')

    conn = get_db_connection()
    pos1 = conn.execute("""SELECT * FROM clients""").fetchall()
    pos2 = conn.execute("""SELECT * FROM things WHERE id_thing NOT IN (SELECT thing_id FROM contracts)""").fetchall()
    pos3 = conn.execute("""SELECT * FROM employees""").fetchall()
    conn.close()
    return render_template('edit_contract.html', contract=contract, clients=pos1, things=pos2, employees=pos3)


@app.route('/send_email/<int:contract_id>', methods=['GET', 'POST'])
def send_email(contract_id):
    contract = get_contract(contract_id)
    if request.method == 'POST':
        email = request.form['email']
        subject = request.form['subject']
        message = request.form['message']
        send_email_to_client(email, subject, message)
        flash('Письмо отправлено успешно!')
        return redirect(url_for('contract', contract_id=contract_id))

    return render_template('send_email.html', contract=contract)


def send_email_to_client(email, subject, message):
    sender_email = "n.remizov00@gmail.com"
    sender_password = "bykn qdgy urpb ghvs"

    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = email
    msg['Subject'] = subject

    msg.attach(MIMEText(message, 'plain'))

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, email, msg.as_string())
        server.close()
    except Exception as e:
        print(f"Failed to send email: {e}")


#biveyqbjtrnrrlxi
















# Опись вещей
#apartment - thing

"""получение всех описей"""
@app.route('/things')
@login_required
def things():
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM contracts, things, clients
    WHERE contracts.thing_id = things.id_thing and contracts.client_id = clients.id_client
    """).fetchall()
    conn.close()
    return render_template('things.html', things=pos)




"""получение одной описи"""
def get_thing(item_id):
    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM contracts, things, clients
    WHERE contracts.thing_id = things.id_thing and contracts.client_id = clients.id_client and things.id_thing = ?
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item

""" Страница-карточка - 1 описи """
@app.route('/thing/<int:thing_id>')
@login_required
def thing(thing_id):
    pos = get_thing(thing_id)
    return render_template('thing.html', thing=pos)




@app.route('/new_thing', methods=('GET', 'POST'))
@login_required
@admin_required
def new_thing():
    """ Страница-добавления новой описи """
    count = None
    price = None
    weight = None
    if request.method == 'POST':
        # добавление новой описи вещей в БД псоле заполнения формы
        try:
            count = int(request.form['count'])
            price = float(request.form['price'])
            insurance = int(request.form['insurance'])
            weight = int(request.form['weight'])
            oversized = int(request.form.get('oversized'))
            client_id = int(request.form.get('owner'))
        except ValueError:
            flash('Некорректные значения')
            client_id = 1
        if not client_id > 0:
            flash('Не все поля заполнены')
        else:
            if not (count and price and weight):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                conn.execute("INSERT INTO 'things' ('count', 'price', 'insurance', 'weight', 'oversized', 'client_id')  VALUES (?, ?, ?, ?, ?, ?)",
                                            (count, price, insurance, weight, oversized, client_id))
                conn.commit()
                conn.close()
                return redirect('/new_contract')

    # отрисовка формы
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM clients""").fetchall()
    conn.close()
    return render_template('new_thing.html', clients=pos)





@app.route('/edit_thing/<int:thing_id>', methods=('GET', 'POST'))
def edit_thing(thing_id):
    thing = get_thing(thing_id)

    if request.method == 'POST':
        try:
            count = int(request.form['count'])
            price = float(request.form['price'])
            insurance = int(request.form['insurance'])
            weight = int(request.form['weight'])
            oversized = int(request.form.get('oversized'))
            client_id = int(request.form.get('owner'))
        except ValueError:
            flash('Некорректные значения')
            client_id = 1
        if not client_id > 0:
            flash('Не все поля заполнены')
        else:
            if not (count and price and weight):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                conn.execute("""
                    UPDATE things
                    SET count = ?, price = ?, insurance = ?, weight = ?, oversized = ?, client_id = ?
                    WHERE id_thing = ?
                """, (count, price, insurance, weight, oversized, client_id, thing_id))

                conn.execute("""
                    UPDATE contracts
                    SET client_id = ?
                    WHERE thing_id = ?
                """, (client_id, thing_id))

                conn.commit()
                conn.close()
                return redirect(f'/thing/{thing_id}')

    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM clients""").fetchall()
    conn.close()
    return render_template('edit_thing.html', thing=thing, clients=pos)


























# Клиенты++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
@app.route('/clients')
@login_required
@admin_required
def clients():
    conn = get_db_connection()
    pos = conn.execute("SELECT * FROM clients").fetchall()
    conn.close()
    return render_template('clients.html', clients=pos)

def get_client(client_id):
    conn = get_db_connection()
    client = conn.execute("SELECT * FROM clients WHERE id_client = ?", (client_id,)).fetchone()
    conn.close()
    if client is None:
        abort(404)
    return client


@app.route('/client/<int:client_id>')
def client(client_id):
    client = get_client(client_id)
    return render_template('client.html', client=client)



@app.route('/new_client', methods=('GET', 'POST'))
@login_required
@admin_required
def new_client():
    if request.method == 'POST':
        name = request.form['name']
        phone_number = request.form['phone_number']
        email = request.form['email']
        passport = request.form['passport']

        if not name:
            flash('Имя обязательно!')
        elif not phone_number:
            flash('Телефон обязателен!')
        elif not email:
            flash('Email обязателен!')
        else:
            conn = get_db_connection()
            conn.execute("INSERT INTO clients (name, phone_number, email, passport) VALUES (?, ?, ?, ?)",
                         (name, phone_number, email, passport))
            conn.commit()
            conn.close()
            return redirect('/clients')

    return render_template('new_client.html')


@app.route('/edit_client/<int:client_id>', methods=('GET', 'POST'))
def edit_client(client_id):
    client = get_client(client_id)

    if request.method == 'POST':
        name = request.form['name']
        phone_number = request.form['phone_number']
        email = request.form['email']
        passport = request.form['passport']

        if not name:
            flash('Имя обязательно!')
        elif not phone_number:
            flash('Телефон обязателен!')
        elif not email:
            flash('Email обязателен!')
        else:
            conn = get_db_connection()
            conn.execute(
                "UPDATE clients SET name = ?, phone_number = ?, email = ?, passport = ? WHERE id_client = ?",
                (name, phone_number, email, passport, client_id)
            )
            conn.commit()
            conn.close()
            return redirect(url_for('client', client_id=client_id))

    return render_template('edit_client.html', client=client)







# Сотрудники++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
@app.route('/employees')
@login_required
@admin_required
def employees():
    conn = get_db_connection()
    pos = conn.execute("SELECT * FROM employees").fetchall()
    conn.close()
    return render_template('employees.html', employees=pos)

def get_employee(employee_id):
    conn = get_db_connection()
    employee = conn.execute("SELECT * FROM employees WHERE id_employee = ?", (employee_id,)).fetchone()
    conn.close()
    if employee is None:
        abort(404)
    return employee

@app.route('/employee/<int:employee_id>')
def employee(employee_id):
    employee = get_employee(employee_id)
    return render_template('employee.html', employee=employee)


@app.route('/new_employee', methods=('GET', 'POST'))
@login_required
@admin_required
def new_employee():
    if request.method == 'POST':
        name = request.form['name']
        position = request.form['position']
        phone_number = request.form['phone_number']
        email = request.form['email']
        department = str(request.form['department'])

        if not name:
            flash('Имя обязательно!')
        elif not position:
            flash('Должность обязательна!')
        elif not phone_number:
            flash('Телефон обязателен!')
        elif not email:
            flash('Email обязателен!')
        else:
            conn = get_db_connection()
            conn.execute("INSERT INTO employees (name, position, phone_number, email, department) VALUES (?, ?, ?, ?, ?)",
                         (name, position, phone_number, email, department))
            conn.commit()
            conn.close()
            return redirect('/employees')

    return render_template('new_employee.html')


@app.route('/edit_employee/<int:employee_id>', methods=('GET', 'POST'))
def edit_employee(employee_id):
    employee = get_employee(employee_id)

    if request.method == 'POST':
        name = request.form['name']
        position = request.form['position']
        phone_number = request.form['phone_number']
        email = request.form['email']
        department = request.form['department']

        if not name:
            flash('Имя обязательно!')
        elif not position:
            flash('Должность обязательна!')
        elif not phone_number:
            flash('Телефон обязателен!')
        elif not email:
            flash('Email обязателен!')
        else:
            conn = get_db_connection()
            conn.execute(
                "UPDATE employees SET name = ?, position = ?, phone_number = ?, email = ?, department = ? WHERE id_employee = ?",
                (name, position, phone_number, email, department, employee_id)
            )
            conn.commit()
            conn.close()
            return redirect(url_for('employee', employee_id=employee_id))

    return render_template('edit_employee.html', employee=employee)


# Функция для получения данных о сотрудниках
def get_all_employees():
    conn = get_db_connection()
    employees = conn.execute("SELECT * FROM employees").fetchall()
    conn.close()
    return employees


@app.route('/current_occupancy_filters')
@login_required
@admin_required
def current_occupancy_filters():
    employees = get_all_employees()
    return render_template('current_occupancy_filters.html', employees=employees)


@app.route('/generate_report', methods=['POST'])
@login_required
@admin_required
def generate_report():
    start_date = request.form.get('start_date')
    end_date = request.form.get('end_date')
    status = request.form.get('status')
    employee_id = request.form.get('employee')

    # Преобразование статуса в числовое значение
    if status == 'completed':
        status_value = 1
    elif status == 'in_progress':
        status_value = 0
    else:
        status_value = 'all'

    query = """SELECT 
                    contracts.number,
                    contracts.date,
                    employees.name,
                    contracts.finish_price,
                    CASE contracts.deal_status
                        WHEN 1 THEN 'Сделка завершена'
                        ELSE 'В работе'
                    END as status,
                    contracts.address_start,
                    contracts.address_finish
                FROM contracts
                LEFT JOIN employees ON contracts.employee_id = employees.id_employee
                WHERE (? IS NULL OR contracts.date >= ?)
                AND (? IS NULL OR contracts.date <= ?)
                AND (? = 'all' OR contracts.deal_status = ?)
                AND (? = 'all' OR contracts.employee_id = ?)"""

    params = [start_date, start_date, end_date, end_date, status_value, status_value, employee_id, employee_id]
    conn = get_db_connection()
    contracts = conn.execute(query, params).fetchall()
    conn.close()

    if not contracts:
        flash('Не найдено контрактов по заданным фильтрам.')

    data = {
        'Номер договора': [contract['number'] for contract in contracts],
        'Дата заключения': [contract['date'] for contract in contracts],
        'Сотрудник': [contract['name'] for contract in contracts],
        'Финальная цена сделки': [contract['finish_price'] for contract in contracts],
        'Статус': [contract['status'] for contract in contracts],
        'Начальный адрес': [contract['address_start'] for contract in contracts],
        'Адрес доставки': [contract['address_finish'] for contract in contracts]
    }

    df = pd.DataFrame(data)

    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Отчет')
    output.seek(0)

    return send_file(output, download_name='current_occupancy_report.xlsx', as_attachment=True)


# Отчеты+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++



@app.route('/reports')
@login_required
@admin_required
def reports():
    conn = get_db_connection()
    pos = conn.execute("SELECT * FROM reports").fetchall()
    conn.close()
    return render_template('reports.html', reports=pos)




def get_report(report_id):
    conn = get_db_connection()
    report = conn.execute("""
        SELECT r.*, e.name as employee_name
        FROM reports r
        LEFT JOIN employees e ON r.employee_id = e.id_employee
        WHERE r.id_report = ?
    """, (report_id,)).fetchone()
    conn.close()
    if report is None:
        abort(404)
    return report


@app.route('/report/<int:report_id>')
def report(report_id):
    report = get_report(report_id)
    return render_template('report.html', report=report)







@app.route('/new_report', methods=('GET', 'POST'))
@login_required
@admin_required
def new_report():

    if request.method == 'POST':
        # добавление нового контракта в БД псоле заполнения формы
        try:
            number = request.form['number']
            date = request.form['date']
            report_type = request.form['report_type']
            description = ""
            employee_id = int(request.form.get('employee'))
        except ValueError:
            flash('Некорректные значения')
            employee_id = 0
        if not (employee_id > 0):
            flash('Не все поля заполнены')
        else:
            if not (number and date):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute("INSERT INTO 'reports' ('number', 'date', 'report_type', 'description', 'employee_id')  VALUES (?, ?, ?, ?, ?)",
                                                (number, date, report_type, description, employee_id))
                conn.commit()
                new_contract_id = cursor.lastrowid
                print(new_contract_id)
                conn.close()
                return redirect('/reports')

    # отрисовка формы
    conn = get_db_connection()
    pos3 = conn.execute("""SELECT * FROM employees""").fetchall()
    conn.close()
    return render_template('new_report.html', employees=pos3)




@app.route('/edit_report/<int:report_id>', methods=('GET', 'POST'))
def edit_report(report_id):
    report = get_report(report_id)

    if request.method == 'POST':
        try:
            number = request.form['number']
            date = request.form['date']
            report_type = request.form['report_type']
            description = str(request.form['description'])
            employee_id = int(request.form.get('employee'))
        except ValueError:
            flash('Некорректные значения')
            employee_id = 0

        if not (employee_id > 0):
            flash('Не все поля заполнены')
        else:
            if not (number and date):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                conn.execute(
                    "UPDATE reports SET number = ?, date = ?, report_type = ?, description = ?, employee_id = ? WHERE id_report = ?",
                    (number, date, report_type, description, employee_id, report_id)
                )
                conn.commit()
                conn.close()
                return redirect(url_for('report', report_id=report_id))

    conn = get_db_connection()
    employees = conn.execute("SELECT * FROM employees").fetchall()
    conn.close()
    return render_template('edit_report.html', report=report, employees=employees)


@app.route('/download_report/<int:report_id>', methods=('GET', 'POST'))
def download_report(report_id):
    report = get_report(report_id)

    if request.method == 'GET' and ('Период:' not in report['description']):
        return render_template('select_period.html', report=report)

    if request.method == 'POST':
        period = request.form['period']

        # Update the report's description with the period
        conn = get_db_connection()
        description = f"Период:{period}"
        conn.execute('UPDATE reports SET description = ? WHERE id_report = ?', (description, report_id))
        conn.commit()
        conn.close()

        report = get_report(report_id)  # Fetch the updated report

        report_type = report['report_type']
        if report_type == "По компании":
            return generate_company_report(report, period)
        elif report_type == "По сотрудникам":
            return generate_employee_report(report, period)

    # Extract the period from the description
    period = report['description'].split('Период:')[1]
    report_type = report['report_type']
    if report_type == "По компании":
        return generate_company_report(report, period)
    elif report_type == "По сотрудникам":
        return generate_employee_report(report, period)


def generate_company_report(report, period):
    conn = get_db_connection()

    template_path = 'reports/report_company.docx'
    doc = Document(template_path)

    # Fill in the template fields
    doc.paragraphs[0].text = doc.paragraphs[0].text.replace('{{NUMBER}}', report['number'])
    doc.paragraphs[3].text = doc.paragraphs[3].text.replace('{{DATE}}', report['date'])
    doc.paragraphs[4].text = doc.paragraphs[4].text.replace('{{EMPLOYEE}}', report['employee_name'])
    doc.paragraphs[7].text = doc.paragraphs[7].text.replace('{{PERIOD}}', period)

    period_parts = period.split(' - ')
    if len(period_parts) != 2:
        raise ValueError("Period must be in 'YYYY-MM-DD - YYYY-MM-DD' format")
    period_start, period_end = period_parts

    contracts = conn.execute("""
        SELECT number, date, address_start, address_finish, finish_price, employee_id
        FROM contracts
        WHERE date BETWEEN ? AND ?
    """, (period_start, period_end)).fetchall()

    table = doc.tables[0]

    total_sum = 0
    for i, contract in enumerate(contracts):
        row = table.add_row().cells
        row[0].text = str(i + 1)
        row[1].text = f"№{contract['number']} от {contract['date']}"
        row[2].text = contract['address_start']
        row[3].text = contract['address_finish']
        row[4].text = str(contract['finish_price'])
        employee = conn.execute("SELECT name FROM employees WHERE id_employee = ?",
                                (contract['employee_id'],)).fetchone()
        row[5].text = employee['name']
        total_sum += contract['finish_price']

    # Fill in the summary table
    doc.tables[1].rows[1].cells[1].text = str(len(contracts))
    doc.tables[1].rows[2].cells[1].text = str(total_sum)

    conn.close()

    # Define the path to save the generated report
    folder_path = os.path.dirname(template_path)
    output_filename = f"generated_report_company_{report['id_report']}.docx"
    output_path = os.path.join(folder_path, output_filename)

    doc.save(output_path)

    return send_file(output_path, as_attachment=True)





@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

if __name__ == '__main__':
    app.run()